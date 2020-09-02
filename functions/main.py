import os

import skimage.io as io
from docx import Document
from docx.shared import Inches
from PIL import Image
from skimage.color import rgb2gray
from functions.text_separation import text_separation
import matplotlib.pyplot as plt
import numpy as np
from functions.thresholding import binarize
from functions.deskew import deskew
from functions.text_segmentation import text_to_lines, lines_to_words, words_to_characters


def main(img_filename):
    img_file_path = os.path.join(os.path.dirname(__file__), '../uploads/' + img_filename)
    img = io.imread(img_file_path)
    img = rgb2gray(img) * 255
    img = text_separation(img)
    img = binarize(img, mode=2)
    lines = text_to_lines(img)
    lines_to_words(lines)
    words_to_characters(lines)

    # img_new_file = Image.fromarray(img, 'F')
    # img_new_file = img_new_file.convert('L')
    # img_new_file.save("your_file.png")

    # document = Document()

    # document.add_heading('Document Title', 0)
    #
    # p = document.add_paragraph('A plain paragraph having some ')
    # p.add_run('bold').bold = True
    # p.add_run(' and some ')
    # p.add_run('italic.').italic = True
    #
    # document.add_heading('Heading, level 1', level=1)
    # document.add_paragraph('Intense quote', style='Intense Quote')
    #
    # document.add_paragraph(
    #     'first item in unordered list', style='List Bullet'
    # )
    # document.add_paragraph(
    #     'first item in ordered list', style='List Number'
    # )

    # document.add_picture('your_file.jpeg', width=Inches(4))

    # records = (
    #     (3, '101', 'Spam'),
    #     (7, '422', 'Eggs'),
    #     (4, '631', 'Spam, spam, eggs, and spam')
    # )
    #
    # table = document.add_table(rows=1, cols=3)
    # hdr_cells = table.rows[0].cells
    # hdr_cells[0].text = 'Qty'
    # hdr_cells[1].text = 'Id'
    # hdr_cells[2].text = 'Desc'
    # for qty, _id, desc in records:
    #     row_cells = table.add_row().cells
    #     row_cells[0].text = str(qty)
    #     row_cells[1].text = _id
    #     row_cells[2].text = desc

    # document.add_page_break()
    #
    # document_name = os.path.splitext(img_filename)[0] + '.docx'
    # document.save(document_name)
    return
