from matplotlib import pyplot as plt
import numpy as np
import skimage.io as io
from docx import Document
import os
from model.model import predict
from docx.shared import Pt


def show_images(images, titles=None):
    n_ims = len(images)
    if titles is None: titles = ['(%d)' % i for i in range(1, n_ims + 1)]
    fig = plt.figure()
    n = 1
    for image, title in zip(images, titles):
        a = fig.add_subplot(1, n_ims, n)
        if image.ndim == 2:
            plt.gray()
        plt.imshow(image)
        a.set_title(title)
        n += 1
    fig.set_size_inches(np.array(fig.get_size_inches()) * n_ims)
    plt.show()


def output_segmentation_results_imgs(lines):
    i = 0
    for line in lines:
        for word in line.words:
            for character in word.characters:
                io.imsave('./results/' + str(i) + '.png', np.array(character.img))
                i += 1


def output_segmentation_results_docx(lines, model, class_names, filename):
    document = Document()
    run = document.add_paragraph().add_run()
    font = run.font
    font.name = 'Calibri'
    font.size = Pt(16)
    for line in lines:
        for word in line.words:
            word_string = ""
            for character in word.characters:
                word_string += predict(model, character.img, class_names)
            run.add_text(word_string + " ")
        run.add_text("\n")
    document_name = os.path.splitext(filename)[0] + '.docx'
    document.save(document_name)
